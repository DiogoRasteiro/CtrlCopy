from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, HttpResponseRedirect
from django.template import loader
from django.core.files.storage import FileSystemStorage
from .models import Document, Student, User
from .forms import DocumentForm, LoginForm
import textract
from nltk.tokenize import sent_tokenize
from difflib import SequenceMatcher
from nltk.corpus import stopwords
from django.contrib.auth import authenticate
from django.contrib.auth import logout as django_logout
from django.contrib.auth import get_user_model
from django.contrib.auth.backends import ModelBackend
from django.contrib.auth import login as auth_login
from collections import Counter
import os
import docx

# Create your views here.
def mainindex(request):
	return render(request, 'mainindex.html')

def index(request):
	document_list = Document.objects.order_by('-id')
	template = loader.get_template('plagiarism/index.html')
	context = {'document_list': document_list}

	return render(request, 'plagiarism/index.html', context)

def login(request):
	# Django Form used for passing the information 
	form = LoginForm()
	if request.method == 'POST': # If information is present in POST, then a user is attempting to login
		username = request.POST.get('username', False)
		password = request.POST['password']
		# Django function for verifying credentials
		user = authenticate(request, username = username, password = password)
		if user is not None:
			auth_login(request, user) # Django function for logging the user in
			request.session['user_type'] = user.user2.user_type
			request.session['upload_stage'] = 0
			return render(request, 'plagiarism/about_us.html')
		else:
			return render(request, 'plagiarism/uploaddoc.html')
	return render(request, 'plagiarism/login.html', {'form':form})

def logout(request):
	django_logout(request)
	return render(request, 'plagiarism/uploaddoc.html')

def about_us(request):
	return render(request, 'plagiarism/about_us.html')

def document(request, document_id):
	document = get_object_or_404(Document, pk = document_id)
	doc_body = list()
	with document.document.open() as var:
		for line in var:
			line = line.rstrip()
			doc_body.append(line)
	return render(request, 'pplagiarismlagiarism/document.html', {'document': document, 'body': str(doc_body)})

def submit(request):
	# This view is separated into 3 stages:
	# 1st is the document submission form
	# 2nd is the document preparation and analysis configuration
	# 3rd is the actual plagiarism detection

	
	if request.method == 'POST' and request.session['upload_stage'] == 0:
		request.session['upload_stage'] = 1
		try:
			doc = Document(author = User.objects.get(username = request.user.get_username()), title = os.path.basename(request.FILES['doc'].name),
			document = request.FILES['doc'], evaluator = User.objects.get(username = 'TestProfessor').user2)
			doc.save()
		except: 
			doc = Document.objects.get(id = request.session['doc_id'])

		request.session['doc_id'] = doc.id
		context = {
			'doc_body' : doc_to_sentences(doc.id),
			'keywords' : list(),
		}

		full_body = ""
		for x in context['doc_body']:
			full_body = full_body + x
		full_body = ' '.join([word for word in full_body.split() if word not in set(stopwords.words('english'))])
		body_split = full_body.split()
		counter = Counter(body_split)
		keywords = counter.most_common(3)
		for key in keywords:
			context['keywords'].append(key[0])

		return render(request, 'plagiarism/AnalyseDoc.html', context)

	if request.method == 'POST' and request.session['upload_stage'] == 1:
		doc = Document.objects.get(id = request.session['doc_id'])

		toCompare = list()
		for doc in list(Document.objects.all()):
			if(doc.id != request.session['doc_id']):
				toCompare.append(doc.id)
		
		results = {
			'plag_sentences': list(),
			'plag_docs': list(),
			'doc_body': doc_to_sentences(doc.id)
		}

		ratio = int(request.POST['percentage']) / 100 # Get Percentage Level defined by user
		plag_sentences_formatted = list() # Auxiliary Array


		for toCompare in toCompare:
			current_plag_sentences = compare_docs(doc_to_sentences(doc.id), doc_to_sentences(toCompare), ratio)
			if len(current_plag_sentences) != 0:
				results['plag_docs'].append(Document.objects.get(id = toCompare).title)
				results['plag_sentences'] = Union(results['plag_sentences'], current_plag_sentences)
				
				for x in current_plag_sentences:
					plag_sentences_formatted.append(" <span title=\"Origin: " + Document.objects.get(id = toCompare).title + "\">" +
				x + " </span> \n")

				results['doc_body'] = [" <span title=\"Origin: " + Document.objects.get(id = toCompare).title + "\">" +
				x + " </span> \n" if x in current_plag_sentences else x for x in results['doc_body']]


		
		totalPlagiarismLevel = round(len(results['plag_sentences']) / len(results['doc_body']) * 100, 2) 
		plagLevel = str(totalPlagiarismLevel) + '%'

		results['plagLevel'] = plagLevel
		results['plag_sentences'] = plag_sentences_formatted
		

		return render(request, 'plagiarism/student_submit.html', results)

	else:
		if  request.session['upload_stage'] != 0 and (request.session['doc_id'] != 0):
			doc = Document.objects.get(id = request.session['doc_id'])
			doc.delete()

		form = DocumentForm()
		request.session['upload_stage'] = 0
		return render(request, 'plagiarism/uploaddoc.html', {'form':form})

def submitted(request):
	doc = Document.objects.get(id = request.session['doc_id'])
	try:
		if request.session['evaluation']:
			doc.approved = True
			doc.save()
	except:
		pass

	request.session['doc_id'] = 0
	return render(request, 'plagiarism/submited.html', {})


def profile(request):
	return render(request, 'plagiarism/profile_student.html')

def welcome(request):
	if request.method == 'POST':
		request.session['upload_stage'] = 0
		try: 
			request.session['doc_id'] = request.POST['doc_id']
		except: 
			request.session['upload_stage'] = 1
			request.session['evaluation'] = True
		return submit(request)


	else:
		doc_list = list()

		if request.user.user2.user_type == "Professor":
			for doc in list(Document.objects.all()):
				if(doc.evaluator == request.user.user2) and not doc.approved:
					doc_list.append(doc)
		context = {
			'doc_list': doc_list,
		}

		return render(request, 'plagiarism/welcome.html', context)


# Auxiliary Methods for running the algorithm

# Joins two lists without repeating the elements
def Union(lst1, lst2): 
    final_list = list(set(lst1) | set(lst2)) 
    return final_list 

# Computes level of similarity between two strings
def similar(str1, str2): 
    a = set(str1.split()) 
    b = set(str2.split())
    c = a.intersection(b)
    return float(len(c)) / (len(a) + len(b) - len(c))

# Takes in a file object and returns a list with its sentences
def doc_to_sentences(doc_id):
	doc = Document.objects.get(id = doc_id )
	docText = ""

	if(os.path.splitext(doc.title)[1] == '.docx'):
		docText = getText(doc.document)
	else:	
		doc_body = list()
		with doc.document.open() as var:
				for line in var:
					line = line.rstrip()
					doc_body.append(line)
		docText = "".join(str(doc_body))
		docText = docText[3:len(docText) - 2]

	return sent_tokenize(docText)

# From the original document and a comparison doc, outputs the original's strings that are plagiarized
def compare_docs(originalSentences, compareSentences, ratio):
	results = list()

	for sentence in originalSentences:
		originalSentence = sentence
		sentence = ' '.join([word for word in sentence.split() if word not in (stopwords.words('english'))])
		for compareSentence in compareSentences:
			compareSentence = ' '.join([word for word in compareSentence.split() if word not in (stopwords.words('english'))])
			currentRatio = round(similar(sentence, compareSentence), 2)
			if(currentRatio > ratio):
				results.append(originalSentence)
				break

	return results

# For non txt files, retrieves the text contents.
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
